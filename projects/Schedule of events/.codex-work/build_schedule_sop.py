from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"C:\kwi-automations\projects\Schedule of events\Schedule Template SOP - Final.docx"
TEMPLATE_PATH = r"C:\kwi-automations\projects\Schedule of events\Schedule of Events - Final.xlsx"

BLUE = "2E74B5"
DARK = "18324A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "C9D2DC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, bold=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        style_run(r, bold=bold)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        style_run(r, size=10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        style_run(r, size=10.5)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK)
        run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 12)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(label + ": ")
    style_run(r, size=10.5, bold=True, color=DARK)
    r = p.add_run(text)
    style_run(r, size=10.5, color="111827")
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_BLUE)
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        style_run(r, size=10, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if i == 0:
                set_cell_shading(cells[i], LIGHT_GRAY)
            p = cells[i].paragraphs[0]
            r = p.add_run(value)
            style_run(r, size=9.5, bold=(i == 0), color="111827")
    doc.add_paragraph()
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name in ["List Bullet", "List Number"]:
    s = styles[style_name]
    s.font.name = "Calibri"
    s.font.size = Pt(10.5)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.25

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("Event Schedule Template SOP")
style_run(r, size=24, bold=True, color=DARK)
subtitle = doc.add_paragraph()
r = subtitle.add_run("Standard operating procedure for the final clean tested Excel template")
style_run(r, size=12, color="55616D")

meta = doc.add_table(rows=4, cols=2)
set_table_width(meta, [1.6, 4.9])
set_table_borders(meta)
meta_data = [
    ("Template file", TEMPLATE_PATH),
    ("Primary owner", "Event operations / schedule owner"),
    ("Version rule", "Save a new version before structural edits or bulk imports."),
    ("Purpose", "Maintain one source of truth for event data and produce readable daily, weekly, location, sponsor, speaker, attendee, and run-of-show views."),
]
for row, (label, value) in zip(meta.rows, meta_data):
    set_cell_shading(row.cells[0], LIGHT_BLUE)
    p = row.cells[0].paragraphs[0]
    style_run(p.add_run(label), size=9.5, bold=True, color=DARK)
    p = row.cells[1].paragraphs[0]
    style_run(p.add_run(value), size=9.5)

add_callout(doc, "Operating rule", "Data Input is the source of truth. All schedule views should be treated as outputs from that table, not as places to type event data.", "FFF3CD")

add_heading(doc, "1. Scope")
add_para(doc, "Use this SOP when creating, updating, testing, printing, or distributing the event schedule workbook. It applies to day-of-event schedules, week-at-a-glance schedules, run-of-show planning, and summary views by location, sponsor, speaker, and required attendee.")

add_heading(doc, "2. Workbook Map")
add_matrix(
    doc,
    ["Sheet", "Primary use", "User action"],
    [
        ("Data Input", "Master event table.", "Enter and edit event rows here."),
        ("Reference", "Dropdown/source lists.", "Add new categories, locations, owners, statuses, and filters inside the table area."),
        ("Daily Schedule", "One-day printout.", "Select date/filter values; do not type into the schedule grid."),
        ("Weekly Calendar", "Formula-driven weekly timed event list.", "Set week start and optional filters; use as the weekly planning printout."),
        ("Run of Show", "Minute-by-minute event rundown.", "Select location and event; fill in cues, leads, tech, and notes."),
        ("Location Schedule", "Printable location utilization list.", "Review rooms/spaces by event time."),
        ("Sponsors / Speakers / Required Attendees", "Printable filtered lists.", "Review only events where those fields are populated."),
        ("Instructions", "Built-in workbook guide.", "Use for quick reminders and onboarding."),
    ],
    [1.45, 2.25, 2.8],
)

add_heading(doc, "3. Adding or Updating Events")
add_numbered(doc, [
    "Open the latest clean template or working copy.",
    "Go to Data Input and add one event per row inside EventInputTable.",
    "Complete Show Name, Organizer, KWI POC, Start Date, Start Time, End Date, End Time, Event Name, Location, Category, Class, and Status.",
    "For timed events, use real Excel time values in Start Time and End Time.",
    "For unknown timing, use TBD. For true all-day items, use All Day.",
    "Enter multiple required attendees, speakers, or sponsors as comma-separated names.",
    "Recalculate or reopen the workbook if a formula-driven view does not update immediately.",
])

add_heading(doc, "4. Adding Categories and Dropdown Values")
add_numbered(doc, [
    "Open Reference.",
    "Add the new plain category name in the Categories column. Do not type the Category: prefix in the source category column.",
    "Use the generated Category: value in Data Input.",
    "Keep new values inside the ReferenceListsTable shaded/table area so validations and formulas continue to pick them up.",
])

add_heading(doc, "5. Reading the Main Views")
add_matrix(
    doc,
    ["View", "What to check", "Common issue"],
    [
        ("Daily Schedule", "Timed events appear in hourly slots; overlaps shift across slot columns.", "If missing, check date/time, filters, location, category, and class."),
        ("TBD row", "TBD events for the selected day appear once.", "If duplicated, verify the bottom table is All-Day Events only."),
        ("All-Day Events", "All-day items for the selected day appear below the timed grid.", "If missing, confirm Start/End Date matches the selected day."),
        ("Weekly Calendar", "Formula-driven timed list for the selected week.", "If blank, check week start, show filter, and numeric start/end times."),
        ("Run of Show", "Minute rows start at selected event start time.", "Requires numeric start and end times."),
    ],
    [1.55, 2.65, 2.3],
)

add_heading(doc, "6. Testing Before Distribution")
add_callout(doc, "Required QA", "Always test on a QA copy, not the clean template. The clean template should remain free of test events.", "E8EEF5")
add_bullets(doc, [
    "Add a new reference category and use it on a test event.",
    "Add three overlapping timed events and verify schedule slots calculate as 1, 2, and 3.",
    "Add a TBD event and verify it appears once in the TBD row.",
    "Add an all-day event and verify it appears once in All-Day Events.",
    "Confirm Weekly Calendar, Location Schedule, Sponsors, Speakers, and Required Attendees update from Data Input.",
    "Select a test event on Run of Show and confirm the minute timeline starts at the correct time.",
    "Save, close, reopen, and re-check the same items.",
])

add_heading(doc, "7. Printing and Distribution")
add_bullets(doc, [
    "Use Daily Schedule for the day-of handout.",
    "Use Weekly Calendar for weekly planning.",
    "Use Location Schedule for room/space coordination.",
    "Use Run of Show for event leads, stage managers, speakers, and AV coordination.",
    "Before printing, check that filters are set correctly and the selected date/week is correct.",
])

add_heading(doc, "8. Version Control")
add_numbered(doc, [
    "Never overwrite the last known-good file.",
    "Use a new version number for structural changes, formula updates, or workflow changes.",
    "Keep a clean template and a separate QA working copy.",
    "After QA passes, distribute only the clean tested template.",
])

doc.add_page_break()
add_heading(doc, "9. Troubleshooting Quick Reference")
add_matrix(
    doc,
    ["Symptom", "Likely cause", "First check"],
    [
        ("Event missing from Daily Schedule", "Date/time/filter mismatch.", "Data Input Start Date, Start Time, End Time, Show, Location, Category, Class."),
        ("Overlap appears in wrong column", "Schedule slot formula or time overlap edge case.", "Rows with same show/date and overlapping numeric times."),
        ("TBD appears in timed grid", "Time value entered inconsistently.", "Use exact TBD in Start/End Time fields."),
        ("All-day event missing", "Start/End Date not matching selected day.", "D/F date fields and All Day text in E/G."),
        ("Summary tab missing new row", "Workbook not recalculated or event field blank.", "Save/reopen and confirm sponsor/speaker/attendee field is populated."),
    ],
    [1.8, 2.2, 2.5],
)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Prepared for the Event Schedule Calendar workbook package - Final")
style_run(r, size=9, color="55616D")

doc.save(OUT)
print(OUT)
